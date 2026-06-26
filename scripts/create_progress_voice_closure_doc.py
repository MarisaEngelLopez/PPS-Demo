from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = ROOT.parent
DOCS_DIR = ROOT / "docs"
OUT = DOCS_DIR / "Project_Ops_System_Project_Progress_Voice_Assistant_Closure_2026-06-03.docx"

BLUE = "1F4E79"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F9"
GREEN = "D9F5E3"
AMBER = "FFF1CC"
WHITE = "FFFFFF"


def latest_backup_name() -> str:
    backups = sorted(
        (WORKSPACE_ROOT / "backups").glob("progress_voice_closure_*"),
        key=lambda path: path.name,
        reverse=True,
    )
    return backups[0].name if backups else "Not found"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=WHITE)
        set_cell_fill(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Cm(widths[i])

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = LIGHT_GREY if row_index % 2 == 0 else WHITE
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_fill(cells[i], fill)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[i].width = Cm(widths[i])

    document.add_paragraph()


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_fill(cell, fill)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)

    document.add_paragraph()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(10 if style_name != "Heading 1" else 14)
        style.paragraph_format.space_after = Pt(5)


def build_document() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Operations System")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Progress Voice Assistant Closure")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {date.today().isoformat()} | Backup: {latest_backup_name()}").italic = True

    add_callout(
        document,
        "Closure Statement",
        "The Project Progress Assistant voice pilot is closed as a working controlled assistant slice. It accepts text or voice input, interprets bounded project progress commands, creates auditable suggestions, and applies changes only after approval.",
        GREEN,
    )

    add_heading(document, "1. Scope Closed", 1)
    add_bullets(
        document,
        [
            "Voice input was added to the Project Progress Assistant using the existing AgentInstruction pipeline with sourceType VOICE.",
            "Text and voice both feed the same bounded interpretation and suggestion workflow.",
            "No direct autonomous project update was introduced; all project changes remain approval controlled.",
            "The assistant remains scoped to the app data universe: selected project, project workstreams, and project milestones.",
        ],
    )

    add_heading(document, "2. Capabilities Implemented", 1)
    add_table(
        document,
        ["Capability", "Behaviour", "Approval impact"],
        [
            ["START_WORKSTREAM", "Sets actualStartDate from the interpreted local date.", "Creates suggestion; approval updates the workstream."],
            ["FINISH_WORKSTREAM", "Sets actualEndDate from the interpreted local date.", "Creates suggestion; approval updates the workstream."],
            ["REOPEN_WORKSTREAM", "Clears actualEndDate.", "Creates suggestion; approval reopens the workstream."],
            ["CHANGE_WORKSTREAM_VISIBILITY", "Supports visibility changes to BOTH, EXECUTIVE, DETAILED, or HIDDEN.", "Creates suggestion; approval updates visibility."],
            ["COMPLETE_EVENT", "Marks the milestone completed and sets the single visible milestone date to today.", "Creates suggestion; approval updates eventDate and completionDate."],
            ["REOPEN_EVENT", "Reopens the milestone by clearing completion state.", "Creates suggestion; approval reopens the milestone."],
            ["CHANGE_EVENT_VISIBILITY", "Supports milestone visibility changes.", "Creates suggestion; approval updates visibility."],
        ],
        [4.5, 7.0, 5.0],
    )

    add_heading(document, "3. Voice And Interpretation Design", 1)
    add_bullets(
        document,
        [
            "Voice capture uses the browser speech recognition surface and leaves the transcript visible for review before interpretation.",
            "The assistant states what it understood before any suggestion is created.",
            "Candidate correction is available even when confidence is medium or high.",
            "Matching is bounded: the interpreter ranks only records that exist in the current app data, not external or open-ended knowledge.",
            "Synonym support was added for Agent/Assistant naming, including PR Agent/PR Assistant and TT Agent/TT Assistant variants.",
            "Common command words such as complete, milestone, event, close, and reopen are excluded from target ranking so they do not overpower the real record name.",
        ],
    )

    add_heading(document, "4. Data And Logging Architecture", 1)
    add_table(
        document,
        ["Element", "Role", "Closure status"],
        [
            ["AgentSourceConfig", "Controls whether PROJECT_PROGRESS voice input is visible and usable.", "VOICE enabled for PROJECT_PROGRESS."],
            ["AgentInstruction", "Stores raw text or voice transcript, source type, parsed intent, project, and target references.", "Used for each interpreted instruction."],
            ["AgentSuggestion", "Stores the proposed project progress change and payload.", "Used for all mutating commands."],
            ["AgentApproval", "Stores approve/reject decisions.", "Required before project data changes."],
            ["AgentActionLog", "Stores instruction, suggestion, approval, and application events.", "Integrated with the central agent transaction log."],
        ],
        [4.1, 8.2, 4.2],
    )

    add_heading(document, "5. Navigation And Menu Closure", 1)
    add_bullets(
        document,
        [
            "The main navigation now groups related functionality under parent menus rather than spreading every assistant as a top-level item.",
            "Organizations are placed before Projects because they are master data used by transactional project records.",
            "Projects includes Projects and Progress Assistant.",
            "Time Tracking includes Time Entries and TT Assistant.",
            "Configuration is a top-level area for agent configuration, transaction logs, and future security/testing configuration.",
            "Dropdown menus now close when the user moves to another menu item, improving speed and avoiding visual clutter.",
            "This navigation structure supports later security segmentation because operational assistants and administrative configuration can be permissioned separately.",
        ],
    )

    add_heading(document, "6. Validation Performed", 1)
    add_table(
        document,
        ["Check", "Result"],
        [
            ["TypeScript", "npx tsc --noEmit passed."],
            ["Lint", "npm run lint passed."],
            ["Route smoke check", "/projects/progress-assistant returned 200 and exposed voice input."],
            ["User functional test", "Voice/text interpretation created suggestions; approval updated project progress records."],
            ["Milestone correction", "Complete milestone now updates the visible milestone date as well as completion state."],
        ],
        [6.0, 10.5],
    )

    add_heading(document, "7. Known Limits And Pragmatic Controls", 1)
    add_bullets(
        document,
        [
            "The assistant is a pilot-level natural language interpreter, not an open-ended LLM planner.",
            "Some rephrasing may still be needed when pronunciation or transcript quality is poor.",
            "Query-style questions such as which workstreams have no actual start date are not yet implemented; they should become a read-only query capability.",
            "No-op and replacement warnings should be added next for cases such as completing an already completed milestone or finishing a workstream that already has an actual end date.",
            "A future rule should handle finishing a workstream with no actual start date as an explicit assisted correction, not as a silent automatic update.",
        ],
    )

    add_heading(document, "8. Recommended Next Steps", 1)
    add_table(
        document,
        ["Priority", "Next step", "Reason"],
        [
            ["1", "Add no-op blockers and replacement warnings for project progress commands.", "Improves data quality without increasing complexity."],
            ["2", "Add read-only Project Progress query capability.", "Useful and low risk because it does not mutate data."],
            ["3", "Document standard voice assistant pattern as reusable architecture.", "Prepares the next agent without duplicating decisions."],
            ["4", "Add automated tests for agent suggestion creation and approval application.", "Protects the controlled mutation pipeline."],
            ["5", "Review mobile layout for assistant panels and navigation.", "Voice usage is expected to happen on the go."],
        ],
        [2.2, 7.2, 7.1],
    )

    add_callout(
        document,
        "Architectural Decision",
        "Voice is an input channel, not a separate agent architecture. The durable pattern remains: Voice or Text -> Transcript/Instruction -> Bounded Interpretation -> Suggestion -> Approval -> Apply -> Log.",
        AMBER,
    )

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
