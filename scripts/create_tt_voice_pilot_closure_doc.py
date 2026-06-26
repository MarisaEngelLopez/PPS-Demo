from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = ROOT.parent
DOCS_DIR = ROOT / "docs"
OUT = DOCS_DIR / "Project_Ops_System_TT_Voice_Natural_Language_Pilot_Closure_2026-06-03.docx"

BLUE = "1F4E79"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F9"
GREEN = "D9F5E3"
AMBER = "FFF1CC"
WHITE = "FFFFFF"


def latest_backup_name() -> str:
    backups = sorted(
        (WORKSPACE_ROOT / "backups").glob("tt_voice_pilot_closure_*"),
        key=lambda path: path.name,
        reverse=True,
    )
    return backups[0].name if backups else "Not found"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=WHITE)
        set_cell_fill(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Cm(widths[i])

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = LIGHT_GREY if row_index % 2 == 0 else WHITE
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_fill(cells[i], fill)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[i].width = Cm(widths[i])

    document.add_paragraph()


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_fill(cell, fill)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    document.add_paragraph()


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Project Ops System - TT Voice/Natural Language Pilot closure - 3 June 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_document() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Project Ops System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("TT Voice / Natural Language Pilot Closure")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Closure document | {date(2026, 6, 3).strftime('%d %B %Y')} | Prepared with Codex")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("6B7280")

    add_callout(
        document,
        "Closure statement",
        "The TT Voice / Natural Language Pilot is frozen as the first operational voice-enabled assistant baseline. It supports controlled text and browser-voice input, transcript review, bounded interpretation inside the application data universe, candidate correction, explicit confirmation, and logged start/pause/resume/finish actions for the Time Tracking Assistant.",
        fill=GREEN,
    )

    add_heading(document, "1. Frozen Scope")
    add_table(
        document,
        ["Capability", "Frozen baseline", "Status"],
        [
            ["Voice source control", "The TIME_TRACKING agent VOICE source is enabled through agent configuration and read by the assistant page.", "Closed"],
            ["Voice capture", "Browser speech recognition captures a transcript. The user can manually stop listening, review, and edit the transcript before interpretation.", "Closed"],
            ["Natural language interpretation", "The interpreter detects start, pause, resume, finish/stop/end/complete instructions and maps them to existing TT workflow actions.", "Closed"],
            ["Candidate correction", "The best workstream match is shown with close candidates. The user can correct the match before confirming.", "Closed"],
            ["Start work session", "Confirmed start creates a WorkSession through the existing TT assistant action and logs source type TEXT or VOICE.", "Closed"],
            ["Pause/resume/finish", "Natural language instructions resolve the one applicable active or paused work session and require confirmation before applying.", "Closed"],
            ["Transactional logging", "AgentInstruction and AgentActionLog records are created; source type and interpretation corrections are preserved in payload/log data.", "Closed"],
        ],
        widths=[4.1, 10.5, 2.5],
    )

    add_heading(document, "2. Interaction Contract")
    add_bullets(
        document,
        [
            "The assistant never applies a voice or natural-language instruction directly from transcript alone.",
            "The user must review the transcript and click Interpret.",
            "The assistant states what it understood in writing.",
            "For start instructions, the user can correct the workstream candidate before confirming.",
            "For pause, resume, and finish, the assistant acts only when exactly one applicable session exists.",
            "If no session or multiple sessions are applicable, the assistant returns a clear message and refuses to guess.",
        ]
    )

    add_heading(document, "3. Architecture Outcomes")
    add_table(
        document,
        ["Layer", "Implementation", "Reason"],
        [
            ["Voice source", "AgentSourceConfig controls whether voice input is available for TIME_TRACKING.", "Keeps behavior administrable and prepares security segmentation."],
            ["Transcript", "Browser speech recognition writes into the same Natural Language Pilot input field.", "Avoids a separate voice architecture."],
            ["Interpreter", "A reusable natural-language interpreter ranks app records and returns structured intent/candidates.", "Keeps the answer universe bounded to existing application data."],
            ["Confirmation", "The existing TT assistant server actions are used after confirmation.", "Reuses hardened transactional logic instead of duplicating write behavior."],
            ["Logging", "Instruction/action logs record source type and correction details.", "Supports auditability and future testing."],
        ],
        widths=[3.4, 8.0, 5.7],
    )

    add_heading(document, "4. Safeguards")
    add_bullets(
        document,
        [
            "Voice input is hidden unless the TIME_TRACKING VOICE source is enabled.",
            "Voice transcripts remain editable before interpretation.",
            "High-confidence matches can still be corrected before confirmation.",
            "Low-confidence start interpretations can be corrected by selecting a valid candidate.",
            "Pause/resume/finish do not search broadly across historical sessions; they resolve only the current applicable work session state.",
            "All actions continue to use existing capability checks and protected TT assistant rules.",
        ]
    )

    add_heading(document, "5. Known Limits Kept Out of Scope")
    add_table(
        document,
        ["Limit", "Current handling", "Future path"],
        [
            ["Pure voice confirmation", "The pilot uses written confirmation and click/tap actions.", "Add spoken confirmation and voice yes/no only after the baseline is stable."],
            ["Browser transcript variability", "Transcript is shown and can be edited before interpretation.", "Add bounded LLM disambiguation only for ambiguous cases if deterministic ranking is insufficient."],
            ["Acronyms and pronunciation", "Candidate correction handles misheard workstream names without creating endless aliases.", "Introduce a configurable alias dictionary only where business value is clear."],
            ["Multiple active sessions", "The assistant refuses to guess.", "Keep as a protected rule unless multi-session support is deliberately introduced."],
            ["Automated tests", "Manual validation completed by user during development.", "Add Automatic Testing Foundation and Agent Testing workstreams later."],
        ],
        widths=[4.1, 7.0, 6.0],
    )

    add_heading(document, "6. Validation Performed")
    add_bullets(
        document,
        [
            "User validated text and voice start instructions.",
            "User validated candidate correction before confirmation.",
            "User validated pause, resume, and finish natural-language instructions.",
            "TypeScript compilation passed with npx tsc --noEmit.",
            "ESLint passed with npm run lint.",
            "The TT Assistant route rendered successfully with voice input enabled.",
        ]
    )

    add_heading(document, "7. Backup Reference")
    add_callout(
        document,
        "Backup created",
        f"Backup folder: {latest_backup_name()}. The backup includes the workspace-root database, the app-local database copy if present, a browsable project-ops-system code copy excluding generated dependency folders, and a compressed code archive.",
        fill=AMBER,
    )

    add_heading(document, "8. Recommended Next Workstreams")
    add_bullets(
        document,
        [
            "Project Progress natural-language pilot using the same bounded candidate/correction architecture.",
            "Automatic Testing Foundation with an isolated test database, fixtures, and baseline commands.",
            "Agent Testing using the foundation to cover TT and Project Progress assistant behaviors.",
            "Optional bounded LLM disambiguation, only after deterministic candidate ranking and correction are measured against real voice usage.",
        ]
    )

    add_footer(document)
    document.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
