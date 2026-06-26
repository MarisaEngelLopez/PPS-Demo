from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT = DOCS_DIR / "Project_Ops_System_Codex_Architecture_Audit_2026-06-01.docx"

BLUE = "1F4E79"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F9"
MID_GREY = "E8EEF5"
GREEN = "D9F5E3"
AMBER = "FFF1CC"
RED = "F8D7DA"
WHITE = "FFFFFF"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=WHITE)
        set_cell_fill(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Cm(widths[i])

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = LIGHT_GREY if row_index % 2 == 0 else WHITE
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_fill(cells[i], fill)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Cm(widths[i])

    document.add_paragraph()


def add_status_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Area", "Rating", "Evidence", "Next action"]
    widths = [4.0, 2.4, 6.2, 4.2]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=WHITE)
        set_cell_fill(cell, BLUE)
        cell.width = Cm(widths[i])
    set_repeat_table_header(table.rows[0])

    rating_fills = {
        "Green": GREEN,
        "Amber": AMBER,
        "Red": RED,
        "Blue": LIGHT_BLUE,
    }

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 1))
            cells[i].width = Cm(widths[i])
            fill = rating_fills.get(value, WHITE) if i == 1 else WHITE
            set_cell_fill(cells[i], fill)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    document.add_paragraph()


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Project Ops System - Codex audit - 1 June 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_document() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Project Ops System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Codex Architecture Stabilization Audit")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Version 1.0 | 1 June 2026 | Post bridge-field cleanup baseline")
    run.font.name = "Calibri"
    run.font.size = Pt(10)

    document.add_paragraph()
    add_table(
        document,
        ["Field", "Value"],
        [
            ["System", "Project Operations System"],
            ["Audit date", "1 June 2026"],
            ["Audit scope", "Architecture, schema, code standardization, reporting outputs, security readiness, and translation readiness"],
            ["Current baseline", "Working Next.js application with Prisma SQLite database, reusable UI primitives, standardized reporting, and bridge fields removed"],
            ["Most recent backup used before bridge cleanup", r"C:\Users\maris\Documents\Project Management system 2\backups\bridge_cleanup_20260601_163318"],
        ],
        widths=[4.0, 12.5],
    )

    add_heading(document, "Executive Summary", 1)
    document.add_paragraph(
        "The system has reached a materially stronger architecture baseline. The main migration bridges have been removed, "
        "the report package is functioning as a single point of truth for executive reporting, and the most important reusable "
        "UI and domain patterns are now in place. The application is not yet production-hardened for security, segmentation, "
        "translation, automated regression testing, or operational audit logging, but it is now in a good state to add those "
        "capabilities progressively."
    )
    add_bullets(
        document,
        [
            "Architecture state: stable development baseline with clean build, clean full lint, valid Prisma schema, and database migrations up to date.",
            "Data state: no active ProjectStatus or RiskStatus bridge tables remain; required project, decision, risk, and risk action status links are populated.",
            "UI state: shared primitives exist for section headers, add buttons, tables, form fields, nested panels, action groups, and cockpit metrics.",
            "Reporting state: executive report screen, PDF, PPT, and Gantt now share a structured reporting model and Gantt output contract.",
            "Main residual risk: production hardening should not start until role-based security, audit logging, automated tests, and translation catalogs are designed.",
        ],
    )

    add_callout(
        document,
        "Audit conclusion",
        "The codebase is now rational enough to continue building safely. The next architectural phase should focus on hardening foundations, not further broad visual refactoring.",
        GREEN,
    )

    add_heading(document, "What Has Been Done Since Starting In Codex", 1)
    add_table(
        document,
        ["Area", "Outcome", "Current status"],
        [
            ["Business code automation", "Automated project, risk, decision, and risk action business code generation. Added configuration tools to review next codes and purge generated test records.", "Working; configuration-level security needed later."],
            ["Project data migration", "Migrated project data into the Prisma database and aligned records to standardized business codes and status structures.", "Working baseline."],
            ["Executive reporting package", "Created and stabilized reporting packs with Draft, Ready, Approved, and Archived lifecycle behavior, version handling, copy-forward logic, read-only approved reports, delete/reset safeguards, and admin reopening.", "Working and user-tested."],
            ["PDF export", "Built landscape executive report output with cover page, report index, narrative sections, decision/risk cockpits, attention sections, and Gantt output.", "Working; Gantt PDF quality confirmed."],
            ["PPT export", "Added PPT download with cover, sections, cockpit outputs, and Gantt slides from the reporting model.", "Working; visual polish remains optional."],
            ["Gantt output contract", "Introduced a shared timeline/Gantt model so screen, PDF, and PPT derive from one data interpretation.", "Key cornerstone completed."],
            ["Cockpit standardization", "Created reusable cockpit contract and consistent lifecycle/attention KPI color semantics across decision, risk, workstream, and milestone reporting.", "Working; future entities can reuse the pattern."],
            ["Business admin standardization", "Standardized statuses, status scopes, status usage, project types, task families, phases, workstreams, events, templates, organizations, contacts, time tracking, and configuration separation.", "Mostly standardized."],
            ["Transactional entities", "Standardized decisions, risks, and risk actions with shared CRUD behavior, status semantics, delete safeguards, action visibility, and expand/collapse controls.", "Working and user-tested."],
            ["Project structure", "Standardized project creation, templates, workstreams, milestones, tasks, subtasks, Gantt visibility, delete rules, and timeline controls.", "Working and user-tested."],
            ["Bridge cleanup", "Removed legacy project manager, sponsor, project status, workstream/task status, risk status, risk action status, and decision status bridges.", "Completed."],
            ["Configuration separation", "Moved high-risk recovery functions out of Admin into Configuration as a future admin-only security boundary.", "Implemented as navigation and route separation; security enforcement still pending."],
        ],
        widths=[3.8, 8.2, 4.5],
    )

    add_heading(document, "Architecture Snapshot", 1)
    add_table(
        document,
        ["Layer", "Current implementation", "Architecture direction"],
        [
            ["App routes", "Next.js app router pages under app/ with server actions for mutations.", "Keep pages thin; move business logic to domain modules before adding security."],
            ["UI components", "Reusable UI primitives in components/ui and entity-specific tables under components/.", "Continue using shared controls for tables, headers, buttons, nested panels, and form fields."],
            ["Domain modules", "lib/domain contains entity rules, validation, query helpers, contracts, reporting models, and cockpit metrics.", "Make this the standard for every entity before expanding complex behavior."],
            ["Reporting model", "Executive reporting uses shared view model, Gantt model, and output contract for screen/PDF/PPT.", "Keep outputs renderer-specific but model-driven."],
            ["Database", "Prisma schema with 38 migrations and SQLite dev database.", "Maintain migration discipline; add backup/restore and release process before production."],
            ["Configuration", "Configuration route separated from Admin for powerful recovery/reset operations.", "Protect with administrator-only role and audit log."],
        ],
        widths=[3.2, 7.0, 6.3],
    )

    add_heading(document, "Audit Scope And Evidence", 1)
    add_table(
        document,
        ["Check", "Result", "Evidence"],
        [
            ["Prisma schema validation", "Pass", "`npx prisma validate` completed successfully."],
            ["Migration status", "Pass", "`npx prisma migrate status` reports 38 migrations and database schema up to date."],
            ["Full lint", "Pass", "`npm run lint -- app components lib scripts` completed clean after correcting one React selector state pattern."],
            ["Production build", "Pass", "`npm run build` completed successfully and generated 34 app routes."],
            ["Bridge tables", "Pass", "Database contains no ProjectStatus or RiskStatus tables."],
            ["Required status links", "Pass", "Projects missing governedStatusId: 0; risks missing statusId: 0; risk actions missing statusId: 0; decisions missing statusId: 0."],
            ["Seed file", "Pass", "`prisma/seed.ts` and the package seed script have been removed."],
            ["Historical references", "Expected", "Old migration files and historical docs still mention bridge fields. This is normal history, not active code dependency."],
        ],
        widths=[4.0, 2.3, 10.2],
    )

    add_heading(document, "Audit Findings", 1)
    add_status_table(
        document,
        [
            ["Bridge-field removal", "Green", "Active schema and code no longer depend on ProjectStatus, RiskStatus, projectManagerId, sponsorId, workstream/task legacy status, risk action text status, or decision status bridges.", "Keep bridge cleanup register as historical evidence."],
            ["Database health", "Green", "Database schema is up to date. Required status references are populated. Business data exists across projects, workstreams, decisions, risks, reporting packs, and time entries.", "Add backup/restore procedure before production."],
            ["Build and lint health", "Green", "Full lint and production build pass. One lint issue in ProjectReportSelector was corrected during audit.", "Add CI checks so this stays automatic."],
            ["Reusable UI adoption", "Amber", "Core shared primitives exist and are widely used. Some older admin/configuration pages still use direct table styles and page-local form markup.", "Convert opportunistically when touching those entities; no urgent broad rewrite needed."],
            ["Domain logic consistency", "Amber", "Decisions, risks, reporting, statuses, templates, projects, workstreams, phases, event types, task families, and project types now have domain modules. Some page-level server actions remain, especially admin/configuration/time tracking.", "Extract shared action factories or entity service helpers where duplication becomes costly."],
            ["Reporting single point of truth", "Green", "Executive screen, PDF, PPT, cockpits, and Gantt use shared report data/model contracts. Gantt PDF is confirmed as high quality.", "Keep PPT visual polish as future enhancement, not blocker."],
            ["Security readiness", "Amber", "Configuration is separated from Admin, and delete safeguards exist. No authentication, role enforcement, row-level segmentation, or audit log yet.", "Design roles before production: administrator, project manager, contributor, viewer."],
            ["Translation readiness", "Amber", "Schema has multilingual fields and report language mode. UI labels, statuses, report text, and validation messages are still mostly hardcoded English strings.", "Introduce a translation catalog and message keys before adding Spanish UI."],
            ["Automated test readiness", "Amber", "Manual user validation has been strong; lint/build cover compilation. No automated regression suite exists for critical workflows.", "Add smoke tests for project creation, report pack lifecycle, decisions, risks/actions, time tracking, and exports."],
            ["Operational safety", "Amber", "Powerful purge/reopen tools exist in Configuration. This is correct structurally but unsafe without roles and audit trail.", "Add confirmation, permissions, and audit records before shared usage."],
            ["Documentation state", "Amber", "Current bridge register is updated. Older context/audit docs are historical and intentionally contain legacy references.", "Add a 'historical' note or archive folder later to avoid confusion."],
        ],
    )

    add_heading(document, "Current Standardization Status By Area", 1)
    add_table(
        document,
        ["Area", "Standardization level", "Notes"],
        [
            ["Executive reporting package", "High", "Reference implementation for domain-driven data, lifecycle safeguards, and output contracts."],
            ["Projects and project execution", "High", "Project, workstream, milestone, task/subtask, timeline, and delete behavior are standardized and user-tested."],
            ["Decisions", "High", "Uses generic Status, scoped usage, delete only when Open, reporting visibility, and executive attention logic."],
            ["Risks and risk actions", "High", "Uses generic Status, mitigation action nesting, delete safeguards, cockpit metrics, and report inclusion logic."],
            ["Organizations and contacts", "Medium-high", "Working and structured around organization records; project manager and sponsor now use contacts."],
            ["Time tracking", "Medium-high", "Functional and standardized enough for current use; page still has direct query/action structure."],
            ["Business admin tables", "Medium-high", "Consistent CRUD, activation/deactivation, safeguards, and table styling. Some reusable form/table primitives can be adopted further over time."],
            ["Configuration tools", "Medium", "Correctly separated for future security; intentionally powerful and should be locked down first."],
            ["PDF output", "High", "Server-side PDF and Gantt output are consistent with screen expectations."],
            ["PPT output", "Medium", "Functionally aligned with shared data model; visual quality can be improved later."],
        ],
        widths=[4.0, 3.2, 9.3],
    )

    add_heading(document, "Recommended Hardening Roadmap", 1)
    add_numbered(
        document,
        [
            "Freeze the current development baseline with a backup and optional source control checkpoint.",
            "Create a small automated regression suite for the workflows that users have validated manually.",
            "Design the security model: roles, route groups, record-level project access, configuration-only permissions, and audit logging.",
            "Implement audit log records for destructive and privileged operations, especially configuration purge and reporting pack reopening.",
            "Introduce a translation catalog and replace hardcoded labels/messages progressively by entity.",
            "Standardize remaining page-level admin/time tracking actions only when changing those areas for business reasons.",
            "Polish PPT output after security and test foundations are in place.",
        ],
    )

    add_heading(document, "Security And Segmentation Readiness", 1)
    document.add_paragraph(
        "The architecture is now ready to receive a security layer, but security is not yet implemented. The most important design choice is to keep "
        "configuration functions separate from normal business administration. This has already started structurally, which is good. The next step is "
        "to enforce it technically."
    )
    add_table(
        document,
        ["Future role", "Typical permissions", "Notes"],
        [
            ["Administrator", "Configuration, status usage, business code reset/purge, reporting pack reopening, user/security setup.", "Highest privilege; all actions should be audit logged."],
            ["Project Manager", "Create/update projects, report packs, decisions, risks, milestones, workstreams, tasks, and time entries for assigned projects.", "Main business role."],
            ["Contributor", "Update assigned risks/actions, decisions, tasks, and time entries.", "Should not control configuration or purge data."],
            ["Viewer", "Read project and executive reporting information.", "Useful for steering committee and client views."],
        ],
        widths=[3.2, 8.2, 5.1],
    )

    add_heading(document, "Translation Readiness", 1)
    add_bullets(
        document,
        [
            "The data model already includes several Spanish fields such as nameEs and descriptionEs.",
            "The report language model exists, but UI and validation strings are still embedded in components/actions.",
            "Statuses should continue to be centralized so labels can be translated once per status rather than per page.",
            "Before adding Spanish UI, introduce a message catalog and helper for validation/action messages.",
            "Avoid translating by duplicating components; translate by passing label dictionaries into shared components.",
        ],
    )

    add_heading(document, "Open Items And Non-Blocking Risks", 1)
    add_table(
        document,
        ["Item", "Risk", "Recommendation"],
        [
            ["Historical docs mention removed bridge fields", "Could confuse future readers.", "Keep them as historical records but add a clear archive note or move to docs/history."],
            ["Configuration purge tools", "Correct for testing but powerful.", "Restrict to administrator role and audit every use."],
            ["PPT visual layout", "Good enough functionally but less polished than PDF.", "Defer until after security/test hardening."],
            ["Mixed styling implementation", "Some pages use direct styles rather than shared components.", "Progressively convert when touching those pages."],
            ["Manual testing reliance", "Regression risk grows with new features.", "Add a focused smoke/regression test suite."],
            ["User model still limited", "Current User supports risk/action owners, not real authentication.", "Redesign identity and membership when implementing security."],
        ],
        widths=[4.2, 5.2, 7.1],
    )

    add_heading(document, "Recommended Next Step", 1)
    add_callout(
        document,
        "Next architectural move",
        "Create a small regression test harness before adding security. The application is now stable enough that tests will protect the standardization work already completed.",
        AMBER,
    )
    document.add_paragraph(
        "After the regression harness is in place, the best next production-level step is the security model: administrator-only Configuration, "
        "project-level access segmentation, audit logging, and role-aware navigation. Translation should follow once the security boundaries and "
        "message catalog pattern are defined."
    )

    add_footer(document)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
