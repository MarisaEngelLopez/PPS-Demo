from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT = DOCS_DIR / "Project_Ops_System_Wave_4_1_Time_Tracking_Assistant_Closure_2026-06-02.docx"

BLUE = "1F4E79"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F9"
MID_GREY = "E8EEF5"
GREEN = "D9F5E3"
AMBER = "FFF1CC"
WHITE = "FFFFFF"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=WHITE)
        set_cell_fill(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Cm(widths[i])

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = LIGHT_GREY if row_index % 2 == 0 else WHITE
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_fill(cells[i], fill)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Cm(widths[i])

    document.add_paragraph()


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_fill(cell, fill)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    document.add_paragraph()


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Project Ops System - Wave 4.1 closure - 2 June 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_document() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Project Ops System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Wave 4.1 - Time Tracking Assistant Closure")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Closure document | 2 June 2026 | Prepared with Codex")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("6B7280")

    add_callout(
        document,
        "Closure statement",
        "Wave 4.1 is functionally closed. The Time Tracking Assistant now provides a controlled one-user assistant model with typed instructions, work sessions, pause/resume intervals, approval before official time entry creation, and transactional logging. The implementation follows the architecture principles agreed for the agent foundation: transparent, configurable, auditable, reusable, and prepared for the future voice pilot.",
        fill=GREEN,
    )

    add_heading(document, "1. Scope Closed")
    add_table(
        document,
        ["Capability", "Implemented result", "Closure status"],
        [
            [
                "Assistant menu structure",
                "The assistant has its own route under Time Tracking, with the transactional Time Tracking page kept clean.",
                "Closed",
            ],
            [
                "Typed instruction model",
                "The first default instruction, 'I am starting work', is available through the instruction template model.",
                "Closed",
            ],
            [
                "Work session lifecycle",
                "Start, pause, resume, finish, cancel, and active-note updates are supported.",
                "Closed",
            ],
            [
                "Time calculation",
                "Pause/resume intervals are accumulated and converted through the configured 15-minute rounding rule.",
                "Closed",
            ],
            [
                "Approval workflow",
                "The assistant creates a suggestion first; the official TimeEntry is created only after approval.",
                "Closed",
            ],
            [
                "Editable approval",
                "Date, hours, workstream, and notes can be corrected before approval.",
                "Closed",
            ],
            [
                "Transactional logging",
                "Agent actions are logged and can be filtered and downloaded as CSV.",
                "Closed",
            ],
        ],
        widths=[4.0, 10.3, 2.8],
    )

    add_heading(document, "2. Architecture Outcomes")
    add_bullets(
        document,
        [
            "The agent foundation remains separated from transactional business records. WorkSession and AgentSuggestion are staging/control objects; TimeEntry remains the official transaction.",
            "Agent behavior is controlled through configuration records, capabilities, instruction templates, and rules rather than hardcoded seeded behavior.",
            "Agent configuration logging and transactional agent action logging are separated, which prepares the system for audit and security review.",
            "The assistant uses shared time-tracking data queries and shared workstream ordering, so manual time tracking and the assistant behave consistently.",
            "The UI uses standard page structure, reusable buttons, tables, field styles, section headers, and responsive layout rules.",
            "Browser-local timestamps are captured from the client and displayed in the configured application time zone, solving the UTC/CET mismatch.",
        ]
    )

    add_heading(document, "3. Data Model Elements")
    add_table(
        document,
        ["Model / element", "Role in Wave 4.1", "Production relevance"],
        [
            [
                "AgentDefinition",
                "Defines the assistant as a configured application capability.",
                "Supports future security, activation, and segmentation.",
            ],
            [
                "AgentCapability",
                "Controls which actions the assistant is allowed to perform.",
                "Avoids hidden behavior and prepares role-based access control.",
            ],
            [
                "AgentInstructionTemplate",
                "Stores reusable instruction prompts such as starting work.",
                "Creates a scalable instruction library for future agents.",
            ],
            [
                "WorkSession",
                "Captures work in progress before it becomes an official time entry.",
                "Keeps active assistant work separate from approved transactions.",
            ],
            [
                "WorkSessionPause",
                "Records pause intervals and supports accumulated duration.",
                "Makes the assistant operational instead of only note-based.",
            ],
            [
                "AgentSuggestion",
                "Stores the proposed time entry that still needs approval.",
                "Preserves human control before data is applied.",
            ],
            [
                "AgentApproval",
                "Records approval outcome for the suggestion.",
                "Creates the audit link between suggestion and applied transaction.",
            ],
            [
                "AgentActionLog",
                "Logs start, pause, resume, finish, approval, rejection, export, and note updates.",
                "Provides downloadable operational evidence.",
            ],
        ],
        widths=[4.0, 6.9, 6.2],
    )

    add_heading(document, "4. Functional Decisions Confirmed")
    add_table(
        document,
        ["Decision", "Confirmed approach"],
        [
            [
                "Assistant location",
                "Dedicated assistant page under Time Tracking, rather than embedding all assistant functions into the transactional table.",
            ],
            [
                "Mobile direction",
                "Assistant layout is constrained and responsive so it can become the pilot pattern for mobile-friendly agent workflows.",
            ],
            [
                "Workstream ordering",
                "Workstreams are ordered by recent use in the last 14 days, then other active, then closed, then inactive. The last used workstream is selected by default.",
            ],
            [
                "Closed/inactive workstreams",
                "They remain selectable for time entry when they belong to the active project, because historical and corrective entries may be needed.",
            ],
            [
                "Time source",
                "The client sends the user's laptop timestamp for operational events; display uses the application time-zone helper.",
            ],
            [
                "Logs",
                "Configuration logs and transactional agent logs remain separate. Transaction logs support filtering and CSV download.",
            ],
        ],
        widths=[4.4, 12.0],
    )

    add_heading(document, "5. Verification Performed")
    add_table(
        document,
        ["Check", "Result"],
        [
            ["Lint", "Passed with npm run lint."],
            ["TypeScript", "Passed with npx tsc --noEmit."],
            ["Production build", "Passed with npm run build using the webpack build path."],
            ["Assistant route", "Time Tracking Assistant loads and exposes the expected assistant actions."],
            ["Transactional route", "Manual Time Tracking remains available and uses the same ordered workstream list."],
            ["Log route", "Transactional log page loads and CSV export returns a CSV file."],
            ["Hydration issue", "Resolved by making time formatting deterministic between server and client."],
            ["User cycle test", "Marisa completed the full assistant cycle successfully."],
        ],
        widths=[5.0, 11.4],
    )

    add_heading(document, "6. Residual Items For The Next Wave")
    add_bullets(
        document,
        [
            "Start the voice pilot as an input channel that feeds the same AgentInstruction -> AgentSuggestion -> Approval -> Apply pipeline.",
            "Decide whether transactional agent logs need their own Configuration-level access restriction once the security model is introduced.",
            "Add automated tests around work-session state transitions, approval conversion, and workstream selection rules.",
            "Continue hardening the capability model before multi-user agents are enabled.",
            "Review whether the TT Assistant should become the reference layout for future agent pages.",
        ]
    )

    add_heading(document, "7. Backup Reference")
    document.add_paragraph(
        "A code and database backup was created at closure time in the parent backups folder. The backup includes a source archive without heavy generated folders and a copy of the active SQLite database used by the app."
    )

    add_footer(document)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
