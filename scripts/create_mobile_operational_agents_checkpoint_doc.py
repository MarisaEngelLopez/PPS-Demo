from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Project_Ops_System_Mobile_Operational_Agents_Checkpoint_2026-06-24.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(71, 85, 105)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F8FAFC"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            if index < len(row.cells):
                set_cell_width(row.cells[index], width)
                row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "Project Operations System | Mobile Operational Agents Checkpoint | 2026-06-24"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Project Operations System")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Mobile Operational Agents Checkpoint")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    metadata = doc.add_table(rows=4, cols=2)
    metadata.style = "Table Grid"
    set_table_geometry(metadata, [2500, 6860])
    rows = [
        ("Date", "2026-06-24"),
        ("Scope", "Summary of implementation completed since the last checkpoint/update."),
        ("Baseline document", "PPS Operational Mobile App Design & Implementation.docx"),
        ("Design principle", "One application and one data model; responsive operational mobile surface, not a separate mobile app."),
    ]
    for row, (label, value) in zip(metadata.rows, rows):
        set_cell_fill(row.cells[0], LIGHT_FILL)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(f"\n{text}")
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_status_table(doc: Document) -> None:
    doc.add_heading("Implemented Capabilities", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 4960, 1900])
    headers = ["Area", "What changed", "Status"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = DARK_BLUE

    rows = [
        (
            "Attention workspace",
            "Added attention items for open or paused time tracking sessions so forgotten sessions surface in the daily operational queue.",
            "Implemented and tested",
        ),
        (
            "TT Agent UX",
            "Split the assistant into Start Work, Open Sessions, and Review Suggestions. Open Sessions now defaults when a session exists, includes voice/text control, and shows intervals inline.",
            "Implemented and tested",
        ),
        (
            "PP Agent UX",
            "Split the Project Progress Agent into Create Suggestion and Review Suggestions, matching the assistant pattern established for TT.",
            "Implemented and tested",
        ),
        (
            "Mobile navigation",
            "Desktop keeps the complete menu. Mobile now exposes only operational entry points: Attention, TT Agent, and PP Agent.",
            "Implemented",
        ),
        (
            "Responsive assistant headers",
            "Agent configuration and agent log links are hidden on mobile; operational context links remain visible where useful.",
            "Implemented",
        ),
        (
            "Agent naming",
            "Standardized visible names to TT Agent, PP Agent, Time Tracking Agent, and Project Progress Agent.",
            "Implemented",
        ),
        (
            "Natural language matching",
            "Aligned PP Agent phase-number matching with TT Agent behavior so references such as phase 7 resolve consistently.",
            "Implemented and tested",
        ),
    ]
    for area, change, status in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = change
        cells[2].text = status


def add_architecture_section(doc: Document) -> None:
    doc.add_heading("Architecture Notes", level=1)
    add_callout(
        doc,
        "Core decision",
        "The mobile experience is not a second application. It is a responsive operational layer over the same routes, actions, database, agent logs, and approval model used by desktop.",
    )

    doc.add_heading("Reusable patterns strengthened", level=2)
    add_bullets(
        doc,
        [
            "Operational action card pattern is now used for agent suggestion and session review workflows.",
            "Assistant pages now use a consistent tab pattern to separate action creation from suggestion review.",
            "Mobile visibility is controlled through responsive classes instead of maintaining duplicated mobile components.",
            "Translation keys were added for new tabs and mobile labels, keeping the multilingual architecture intact.",
        ]
    )

    doc.add_heading("Pragmatic constraints preserved", level=2)
    add_bullets(
        doc,
        [
            "Routes remain available even when not shown in the mobile menu, so attention links can still open risks, projects, decisions, or reporting records.",
            "Desktop remains a full management surface; mobile is intentionally limited to operational execution.",
            "Agent audit logs and configuration remain desktop/admin-oriented, not primary mobile workflows.",
        ]
    )


def add_testing_section(doc: Document) -> None:
    doc.add_heading("Validation and Current Behavior", level=1)
    add_bullets(
        doc,
        [
            "TypeScript and lint checks passed after each implementation batch.",
            "User testing confirmed TT Agent and PP Agent tabs work on mobile and laptop.",
            "User testing confirmed PP Agent phase 7 matching works after alignment with TT Agent logic.",
            "User testing confirmed mobile buttons and Cloudflare access are functional after previous server-action origin configuration.",
            "A menu issue where both desktop and mobile menus were visible was corrected by moving display behavior from inline styles into responsive CSS classes.",
        ]
    )


def add_next_steps(doc: Document) -> None:
    doc.add_heading("Recommended Next Steps", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 5560, 1900])
    headers = ["Priority", "Next step", "Reason"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = DARK_BLUE

    rows = [
        (
            "High",
            "Continue mobile user testing with real operational scenarios.",
            "The mobile layer is deliberately narrow; real use will reveal the next friction points faster than adding features speculatively.",
        ),
        (
            "High",
            "Decide whether voice confirmation should speak back the interpreted action.",
            "This is likely the next meaningful step toward true on-the-go operation.",
        ),
        (
            "Medium",
            "Apply the assistant tab/card pattern to future agents as they are created.",
            "This keeps new agent workflows consistent without creating a heavy framework too early.",
        ),
        (
            "Medium",
            "Document mobile-ready versus desktop-only routes as part of future security design.",
            "This prepares the later profile/security layer without creating rework.",
        ),
    ]
    for priority, step, reason in rows:
        cells = table.add_row().cells
        cells[0].text = priority
        cells[1].text = step
        cells[2].text = reason


def build() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Since the last checkpoint, the application has moved from having agents embedded inside broad desktop pages toward a clearer operational assistant model. The key result is a single application that behaves differently by context: full administration and management on desktop, and a focused operational surface on mobile."
    )
    doc.add_paragraph(
        "The main implemented shift is that mobile navigation now concentrates on daily attention and the two active agents, while contextual links still allow the user to reach the relevant underlying records when action is required."
    )

    add_status_table(doc)
    add_architecture_section(doc)
    add_testing_section(doc)
    add_next_steps(doc)

    doc.add_heading("Files and Technical Touchpoints", level=1)
    add_bullets(
        doc,
        [
            "components/ui/MainNav.tsx - responsive desktop/mobile navigation surfaces.",
            "app/globals.css - mobile operational nav and mobile-hidden responsive classes.",
            "components/agents/TimeTrackingAssistant.tsx - three-tab TT Agent operational flow.",
            "components/agents/ProjectProgressAssistant.tsx - two-tab PP Agent flow.",
            "lib/domain/attention/attentionEngine.ts - open work session attention items.",
            "app/projects/progress-assistant-actions.ts - phase-number matching consistency.",
            "lib/i18n/dictionaries.ts - labels for agent naming, tabs, and mobile navigation.",
        ]
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
