from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUTPUT = "docs/Project_Ops_System_Executive_Risk_Review_Package_Checkpoint_2026-06-10.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "1F2937",
    "muted": "64748B",
    "header_fill": "F2F4F7",
    "callout": "F8FAFC",
    "amber": "FFFBEB",
    "border": "D9E2EC",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"]):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    styles = doc.styles
    if "Checkpoint Callout" not in styles:
        callout = styles.add_style("Checkpoint Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Calibri"
        callout.font.size = Pt(10)
        callout.font.color.rgb = RGBColor(31, 41, 55)
        callout.paragraph_format.space_before = Pt(4)
        callout.paragraph_format.space_after = Pt(8)
        callout.paragraph_format.left_indent = Inches(0.12)
        callout.paragraph_format.right_indent = Inches(0.12)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Project Operations System")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Executive Risk Review Package and Architecture Checkpoint")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["blue"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Checkpoint date: 10 June 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_callout(doc, text, fill=COLORS["callout"]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table, 100, 160, 100, 160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    p.style = doc.styles["Checkpoint Callout"]
    p.add_run(text).bold = True
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, COLORS["header_fill"])
        set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    doc.add_paragraph()
    return table


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Project Operations System - Architecture checkpoint")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def build():
    doc = Document()
    style_doc(doc)
    add_footer(doc)
    add_title(doc)

    add_callout(
        doc,
        "Purpose: document the functional and architectural work completed since the last saved design documents, with emphasis on the executive risk review package, single cockpit truth, Gantt alignment, and agent/mobile learning.",
    )

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The system has moved from isolated functional implementation toward a more coherent product architecture. The latest work focused on making the executive report package trustworthy: risk, decision, cockpit, Gantt, and management-review outputs now derive from shared contracts and the same lifecycle semantics used in the transactional application."
    )
    doc.add_paragraph(
        "The most important principle reinforced in this wave is that management-facing outputs must not become a second version of the truth. Screen, PDF, PPT, and data export may use different layouts, but they must share metric definitions, filtering rules, lifecycle logic, and evidence context."
    )

    doc.add_heading("2. Main Outcomes Since The Last Saved Documents", level=1)
    add_table(
        doc,
        ["Area", "Outcome", "Architectural Meaning"],
        [
            [
                "Attention Workspace",
                "Daily attention items became actionable links back to source records.",
                "Derived attention is read-only and traceable, avoiding duplicate operational data.",
            ],
            [
                "Mobile testing",
                "Cloudflare tunnel enabled mobile access; touch and readability issues were identified and partially corrected.",
                "Mobile is treated as a usability target, not a separate product architecture.",
            ],
            [
                "Agents",
                "Time Tracking and Project Progress assistants were expanded with natural language and voice pilots.",
                "Agent instructions, suggestions, approvals, and logs remain controlled and auditable.",
            ],
            [
                "Risk lifecycle",
                "Evidence types, structured evidence, assessments, reviews, review types, review outcomes, and closure safeguards were implemented.",
                "Risk closure is governed by explicit lifecycle evidence, not manual status alone.",
            ],
            [
                "Executive reporting",
                "Cockpits and Gantt outputs were aligned across screen, PDF, PPT, and data export.",
                "Report outputs now consume shared contracts rather than maintaining local variants.",
            ],
        ],
        [1700, 3800, 3860],
    )

    doc.add_heading("3. Executive Risk Review Package", level=1)
    doc.add_paragraph(
        "The executive risk review package now follows a deliberate sequence designed for executive consumption: cockpit, attention, lifecycle summary, and one management review detail page for each risk pending management review."
    )
    add_table(
        doc,
        ["Section", "Purpose", "Design Choice"],
        [
            [
                "Risk Cockpit",
                "Show lifecycle and attention KPIs.",
                "White background because KPI colors already carry meaning.",
            ],
            [
                "Risk Attention",
                "Show exceptions requiring executive awareness.",
                "Light warm background to draw attention without becoming visually aggressive.",
            ],
            [
                "Risk Lifecycle Summary",
                "Show the complete control table ordered by lifecycle urgency.",
                "Light neutral background for governance/control-table reading.",
            ],
            [
                "Management Review Detail",
                "Show the basis for management review and decision-making.",
                "Light amber page with white subsection panels and a single amber left accent.",
            ],
        ],
        [2100, 3800, 3460],
    )

    doc.add_heading("4. Management Review Detail Page", level=1)
    doc.add_paragraph(
        "The management review detail page was enriched so executives can understand the basis of the review without opening the transactional risk record. It remains compact and avoids dumping all operational text into the report."
    )
    add_table(
        doc,
        ["Subsection", "Content Included"],
        [
            ["Risk Snapshot", "Risk code/title, owner, category, status, initial exposure, residual exposure, target date, latest review outcome, and description."],
            ["Assessment Basis", "Latest inherent assessment and latest residual assessment, including probability, impact, exposure, date, assessor, and comments."],
            ["Mitigation & Evidence", "Mitigation actions plus an evidence summary showing type, title, date, reference, and related action."],
            ["Review Context", "Latest review type, outcome, reviewer, date, and comments."],
            ["Linked Decisions", "Decision references linked to the risk review."],
        ],
        [2400, 6960],
    )
    add_callout(
        doc,
        "Design decision: evidence is summarized as type, title, date, reference, and action. Full evidence descriptions stay in the app to preserve executive readability and avoid overwhelming the report.",
        COLORS["amber"],
    )

    doc.add_heading("5. Single Source Of Truth Corrections", level=1)
    doc.add_paragraph(
        "Several discrepancies were detected through testing. They were addressed by moving logic into shared contracts/adapters instead of copying visual or metric definitions into each output."
    )
    add_table(
        doc,
        ["Issue Detected", "Correction"],
        [
            [
                "Risk and decision cockpits differed between transactional pages and executive package.",
                "Executive report now consumes shared cockpit metric contracts and status-usage semantics.",
            ],
            [
                "Executive Gantt counted inactive workstreams and created false phase delays.",
                "Inactive workstreams and linked milestones are excluded from cockpit counts, Gantt bounds, phase summaries, and report timeline rows.",
            ],
            [
                "Workstream lifecycle mixed overdue/delayed with lifecycle states.",
                "Lifecycle is now Open, In Progress, Completed; overdue is treated as attention/variance.",
            ],
            [
                "Risk review package lacked assessment/evidence context.",
                "Executive risk query now includes assessments, reviews, review outcomes, linked decisions, and evidence records.",
            ],
        ],
        [3600, 5760],
    )

    doc.add_heading("6. Architecture Principles Reinforced", level=1)
    add_table(
        doc,
        ["Principle", "Current Application"],
        [
            [
                "Contract-first reporting",
                "Gantt, cockpit, and risk lifecycle outputs use shared contracts/adapters before renderer-specific layout.",
            ],
            [
                "Renderer-specific layout only",
                "Screen, PDF, and PPT may have different geometry, but not different business logic.",
            ],
            [
                "Configurable semantics",
                "Status usage, review types, review outcomes, and evidence types are controlled through admin/configuration data.",
            ],
            [
                "No black-box agents",
                "Agent suggestions require confirmation/approval and produce logs suitable for administration and future audit.",
            ],
            [
                "Progressive hardening",
                "New entities are standardized in phases and verified after each step to protect existing functionality.",
            ],
        ],
        [2700, 6660],
    )

    doc.add_heading("7. Verification Completed", level=1)
    add_table(
        doc,
        ["Check", "Result"],
        [
            ["TypeScript", "npx tsc --noEmit passed."],
            ["Lint", "npm run lint passed."],
            ["Production build", "npm run build passed after the executive risk package changes."],
            ["User functional testing", "Risk lifecycle, assessments, management review trigger, executive package, Gantt alignment, and agent flows have been tested progressively by the user."],
        ],
        [2600, 6760],
    )

    doc.add_heading("8. Recommended Next Steps", level=1)
    add_table(
        doc,
        ["Priority", "Next Step", "Reason"],
        [
            [
                "1",
                "User test the enriched executive risk review package in screen, PDF, and PPT.",
                "This confirms the management-review detail is useful before adding further entities.",
            ],
            [
                "2",
                "Add risk lifecycle items to the Daily Attention workspace.",
                "Management-review-needed risks should surface in daily operations, not only reports.",
            ],
            [
                "3",
                "Begin automatic testing foundation with separate test database.",
                "Protects real project data while stabilizing regression testing entity by entity.",
            ],
            [
                "4",
                "Continue mobile usability review.",
                "Voice and assistant workflows have a different ergonomic profile on mobile.",
            ],
            [
                "5",
                "Document security and configuration separation before user model implementation.",
                "Configuration/admin functions already anticipate future role-based access.",
            ],
        ],
        [900, 4100, 4360],
    )

    doc.add_heading("9. Current Status", level=1)
    doc.add_paragraph(
        "The executive risk review package is now structurally complete for the current risk lifecycle design. The remaining work is user validation, polish of dense PPT output if needed, and integration of lifecycle attention into the home attention workspace."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
