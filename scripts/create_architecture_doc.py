from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT = DOCS_DIR / "Project_Ops_System_Architecture_Hardening_Plan_2026-05-28.docx"

BACKUP_DIR = Path(
    r"C:\Users\maris\Documents\Project Management system 2\backups"
    r"\reporting_architecture_complete_20260528_225315"
)


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F9"
GREEN = "D9F5E3"
AMBER = "FFF1CC"
RED = "F8D7DA"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="FFFFFF")
        set_cell_fill(hdr[i], BLUE)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Cm(widths[i])
        for cell in cells:
            set_cell_fill(cell, LIGHT_GREY if len(table.rows) % 2 == 0 else "FFFFFF")
    document.add_paragraph()
    return table


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE)
        run.font.name = "Calibri"


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)


def build_document() -> None:
    DOCS_DIR.mkdir(exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Ops System")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Architecture Stabilization and Hardening Plan")
    run.bold = True
    run.font.size = Pt(16)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version 1.0 | 28 May 2026 | Reporting package baseline completed")

    document.add_paragraph()
    add_table(
        document,
        ["Field", "Value"],
        [
            ["System", "Project Operations System"],
            ["Architecture status", "Working application with first standardized domain slice completed"],
            ["Current focus", "Reporting package stabilized; next recommended focus is admin tables"],
            ["Database backup", str(BACKUP_DIR / "dev_20260528_225315.db")],
            ["Code backup", str(BACKUP_DIR / "project-ops-system-source_20260528_225315.zip")],
        ],
        widths=[4.5, 11.5],
    )

    add_heading(document, "Executive Summary", 1)
    document.add_paragraph(
        "The application is now functioning as a single point of truth for the project reporting package. "
        "The executive report package has been stabilized across the application workflow, data processing, "
        "PDF export, PPT export, and the Gantt output contract. The recommended architectural direction is "
        "to continue standardizing progressively by entity, preserving the currently working behavior while "
        "moving business rules out of pages and into reusable domain modules."
    )
    add_bullets(
        document,
        [
            "The reporting package is the first completed standardization slice and should now be treated as the reference pattern.",
            "The next best step is to standardize admin tables first, because they are lower risk and define reusable CRUD behavior.",
            "After admin tables, transactional entities should be standardized in priority order: Decisions, Risks, Workstreams/Events, and Time Tracking.",
            "Security, segmentation, auditability, and translation should be built on top of the standardized structure rather than added into legacy patterns.",
        ],
    )

    add_heading(document, "What Has Been Stabilized So Far", 1)
    add_table(
        document,
        ["Area", "Completed work", "Current status"],
        [
            [
                "Business code automation",
                "Automated generation and migration of project business codes.",
                "Working baseline; available for future standardization.",
            ],
            [
                "Project data migration",
                "Project data migrated into the application database and validated through usage.",
                "Working baseline.",
            ],
            [
                "Reporting package lifecycle",
                "Draft, Ready, Approved, and Archived status behavior. Draft deletion and admin rollback were added for controlled testing and correction.",
                "Working and user-tested.",
            ],
            [
                "Executive report PDF",
                "Cover page, index, narrative sections, cockpits, decisions, risks, and Gantt export standardized around the reporting package.",
                "Working; Gantt PDF quality confirmed as correct.",
            ],
            [
                "Executive report PPT",
                "PPT download added with shared data model and Gantt details. Structural parity achieved; visual polish can remain a future enhancement.",
                "Working; future polish optional.",
            ],
            [
                "Gantt output",
                "Introduced a shared Gantt model and output contract so screen, PDF, and PPT no longer depend on separate interpretations.",
                "Working; cornerstone for report consistency.",
            ],
        ],
        widths=[3.5, 8.5, 4.0],
    )

    add_heading(document, "Current Architecture State", 1)
    document.add_paragraph(
        "The application is a Next.js application using Prisma with a SQLite development database. "
        "Routes, UI components, domain services, report renderers, and database definitions are already separated in several places, "
        "but this separation is not yet consistent across all entities. The reporting package now provides the clearest reference pattern."
    )
    add_table(
        document,
        ["Layer", "Current role", "Target direction"],
        [
            ["Routes", "Application pages and API routes under app/.", "Keep thin; delegate business decisions to domain modules."],
            ["Components", "Reusable UI under components/, with reporting components already split.", "Create common entity UI patterns for tables, forms, filters, and dialogs."],
            ["Domain logic", "Reporting domain modules exist under lib/domain/reporting/.", "Move each entity to the same domain-first pattern."],
            ["Export logic", "PDF and PPT now use shared report data and Gantt contracts.", "Use model -> output contract -> renderer for all exportable views."],
            ["Database", "Prisma schema and SQLite database in development.", "Add stronger migration discipline and backup/restore procedures before production."],
        ],
        widths=[3.2, 6.2, 6.2],
    )

    add_heading(document, "Reporting Package Architecture Decisions", 1)
    add_bullets(
        document,
        [
            "The reporting package owns the executive reporting version lifecycle.",
            "Ready reports remain editable for typo correction and final cleanup.",
            "Approved reports become locked to preserve the record, with admin override available when testing or controlled correction is required.",
            "New drafts copy from the latest Ready or Approved report so the next reporting cycle starts from the last controlled version.",
            "The Gantt is now treated as a shared output contract rather than separate screen, PDF, and PPT logic.",
            "Risks and mitigation actions, decisions and decision details, and last-period decisions are grouped on their executive pages according to reporting purpose.",
        ],
    )

    add_heading(document, "Target Production-Level Structure", 1)
    document.add_paragraph(
        "The target structure should make every entity predictable: routes orchestrate, domain modules decide, reusable components display, "
        "and export renderers consume explicit output contracts."
    )
    add_table(
        document,
        ["Folder pattern", "Purpose"],
        [
            ["lib/domain/<entity>/<entity>Types.ts", "Canonical TypeScript types and normalized enums for the entity."],
            ["lib/domain/<entity>/<entity>Queries.ts", "Database read models and include/select patterns."],
            ["lib/domain/<entity>/<entity>Rules.ts", "Business rules, status transitions, visibility logic, and derived flags."],
            ["lib/domain/<entity>/<entity>Validation.ts", "Create/update validation and duplicate protection."],
            ["lib/domain/<entity>/<entity>Metrics.ts", "Cockpit metrics and derived dashboard counters when needed."],
            ["components/<entity>/", "Reusable entity table, editor, filters, status chips, and detail views."],
            ["app/<entity>/page.tsx", "Thin page composition only."],
            ["app/<entity>/actions.ts", "Server action orchestration, calling validation and rules."],
            ["lib/reporting/", "Output contracts and renderers for PDF, PPT, and future formats."],
        ],
        widths=[6.2, 9.4],
    )

    add_heading(document, "Standard Entity Behavior", 1)
    document.add_paragraph(
        "Each entity page should behave consistently unless there is an intentional business exception. "
        "This is the foundation for safe training, future security, and translation."
    )
    add_table(
        document,
        ["Capability", "Standard expectation"],
        [
            ["Create", "Same form behavior, validation feedback, required fields, and duplicate handling."],
            ["Update", "Same edit pattern, optimistic user feedback where appropriate, and no hidden business rule duplication."],
            ["Delete / Archive", "Consistent delete guard. Use archive when records are referenced by reports, transactions, or audit history."],
            ["Status changes", "Defined status transitions in rules files, not scattered through components."],
            ["Visibility", "One shared rule for whether an item is shown transactionally, in reporting, or both."],
            ["Tables", "Consistent sorting, filtering, empty states, row actions, and status chips."],
            ["Validation", "Shared validation per entity, used by server actions and future APIs."],
            ["Exports", "Exportable views consume a stable output contract rather than re-reading raw UI state."],
        ],
        widths=[4.0, 11.8],
    )

    add_heading(document, "Recommended Implementation Sequence", 1)
    add_numbered(
        document,
        [
            "Freeze the reporting package as the reference implementation, allowing only focused bug fixes and later visual polish.",
            "Standardize admin tables first: phases, workstreams, event types, statuses, status scopes/usages, project types, task families, business codes, and reporting pack administration.",
            "Create reusable admin components and domain helpers from the admin-table pilot.",
            "Standardize Decisions next, because the inconsistency between the transactional decision cockpit and reporting has already shown why single-point-of-truth rules matter.",
            "Standardize Risks and mitigation actions after Decisions, including shared reporting and cockpit logic.",
            "Standardize Workstreams, events, milestones, and the Gantt-related transactional model.",
            "Move Time Tracking out of the wrong admin-style structure and into its own transactional domain.",
            "Add production hardening: authentication, roles, organization/project segmentation, audit trail, translation, logging, tests, and backup/restore procedures.",
        ],
    )

    add_heading(document, "Admin Tables First: Rationale", 1)
    document.add_paragraph(
        "Admin tables are the right next area because they are foundational, repeated often, and lower risk than transactional entities. "
        "They will let us create and test the reusable table, form, validation, and delete/archive patterns before applying them to higher-value records."
    )
    add_table(
        document,
        ["Benefit", "Why it matters"],
        [
            ["Low blast radius", "Errors are easier to detect and correct than in transactional reporting records."],
            ["Pattern creation", "Admin entities expose the recurring CRUD behaviors that should become reusable components."],
            ["Future security", "Admin tables are the natural starting point for role restrictions and segmentation."],
            ["Translation readiness", "Admin labels, statuses, and controlled vocabularies are where translation should begin."],
            ["Transactional consistency", "Decisions and risks will later consume standardized statuses and controlled values instead of redefining them."],
        ],
        widths=[4.0, 11.8],
    )

    add_heading(document, "Reusable Elements To Create", 1)
    add_bullets(
        document,
        [
            "EntityPageShell for consistent page structure, title area, actions, and content layout.",
            "EntityToolbar for search, filters, create actions, and context actions.",
            "EntityTable for consistent columns, row actions, empty states, and loading states.",
            "EntityEditor or EntityFormDialog for create and update behavior.",
            "StatusChip, VisibilityChip, and OwnerChip for consistent visual language.",
            "ConfirmDelete or ConfirmArchive for protected destructive actions.",
            "Domain result and validation helpers for predictable server-action responses.",
            "Date, status, and enum formatting utilities with translation keys prepared for later internationalization.",
            "Output contract helpers for PDF, PPT, and future document exports.",
        ],
    )

    add_heading(document, "Security, Segmentation, and Translation Path", 1)
    add_table(
        document,
        ["Capability", "Architecture implication"],
        [
            ["Authentication", "Introduce users and sessions without mixing identity checks into every component."],
            ["Authorization", "Define roles and permissions centrally, then enforce them in server actions and queries."],
            ["Segmentation", "Scope all queries by organization and project membership before production use."],
            ["Audit trail", "Record who changed status, data, and report versions; preserve approved report history."],
            ["Translation", "Move display text, statuses, and controlled labels to translation dictionaries and stable keys."],
            ["Testing", "Add rule-level tests first, then page and export smoke tests for high-value workflows."],
        ],
        widths=[4.0, 11.8],
    )

    add_heading(document, "Definition of Done For Each Standardized Entity", 1)
    add_bullets(
        document,
        [
            "Business rules live in the entity domain module and are reused by screen, reporting, and exports.",
            "Create, update, delete/archive, status changes, validation, and duplicate handling are implemented consistently.",
            "The entity page uses standard reusable UI components unless a documented exception is required.",
            "Cockpit metrics and report visibility are derived from one shared rule set.",
            "Build and focused tests or smoke checks pass before moving to the next entity.",
            "Any required database migration is documented and backed up before execution.",
        ],
    )

    add_heading(document, "Backup Record", 1)
    document.add_paragraph(
        "A code and database backup was created before continuing architecture hardening. "
        "This creates a restore point for the completed reporting-package baseline."
    )
    add_table(
        document,
        ["Backup item", "Path"],
        [
            ["Backup folder", str(BACKUP_DIR)],
            ["Source code zip", str(BACKUP_DIR / "project-ops-system-source_20260528_225315.zip")],
            ["SQLite database copy", str(BACKUP_DIR / "dev_20260528_225315.db")],
        ],
        widths=[4.0, 11.8],
    )

    add_heading(document, "Immediate Next Step", 1)
    document.add_paragraph(
        "Begin the admin-table standardization pilot. The first pilot should define the reusable admin entity contract, "
        "apply it to one simple admin table, verify behavior in the application, and then extend the pattern progressively to the remaining admin tables."
    )

    document.core_properties.title = "Project Ops System Architecture Stabilization and Hardening Plan"
    document.core_properties.subject = "Architecture documentation and next-step hardening plan"
    document.core_properties.author = "Project Operations System"
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
