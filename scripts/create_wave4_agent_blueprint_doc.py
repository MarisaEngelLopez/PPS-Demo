from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT = DOCS_DIR / "Project_Ops_System_Wave_4_One_User_Agent_Foundation_Blueprint_v1_1_2026-06-01.docx"

BLUE = "1F4E79"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F9"
MID_GREY = "E8EEF5"
GREEN = "D9F5E3"
AMBER = "FFF1CC"
WHITE = "FFFFFF"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=WHITE)
        set_cell_fill(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Cm(widths[i])

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = LIGHT_GREY if row_index % 2 == 0 else WHITE
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_fill(cells[i], fill)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Cm(widths[i])

    document.add_paragraph()


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_fill(cell, fill)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    document.add_paragraph()


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Project Ops System - Wave 4 blueprint - 1 June 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_document() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Project Ops System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Wave 4 - One-User Agent Foundation and Time Tracking Assistant")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Blueprint v1.1 | 1 June 2026 | Prepared with Codex")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("6B7280")

    add_callout(
        document,
        "Purpose",
        "This document defines the agreed architecture and implementation contract for Wave 4 before the Time Tracking Assistant is built. The foundation is designed to avoid black-box behavior: every instruction, suggestion, approval, applied change, and work-session event must be visible and auditable.",
    )

    add_heading(document, "1. Executive Summary")
    document.add_paragraph(
        "Wave 4 introduces the first agent-ready layer of the Project Ops System. The objective is not to create an autonomous agent that changes project data directly. The objective is to create a controlled assistant that can capture instructions, produce structured suggestions, require approval, and log what happened."
    )
    add_bullets(
        document,
        [
            "Initial scope: one-user mode for Marisa, focused on the Time Tracking Assistant.",
            "Input channels: TEXT is implemented first; VOICE is designed into the model and piloted early through transcript review.",
            "Execution principle: no agent suggestion is applied to business data without approval.",
            "Business data remains clean: official TimeEntry records are created only after approved suggestions.",
            "Operational layer: WorkSession and WorkSessionPause capture active work before it becomes an official time entry.",
        ],
    )

    add_heading(document, "2. Architecture Principles")
    add_table(
        document,
        ["Principle", "Meaning", "Design consequence"],
        [
            [
                "No black box",
                "The system must show what was asked, what was understood, what was suggested, who approved it, and what changed.",
                "AgentInstruction, AgentSuggestion, AgentApproval, and AgentActionLog are mandatory foundation records.",
            ],
            [
                "Approval before apply",
                "Agent output is a proposal first, not a direct database write.",
                "Agent suggestions must be approved, edited, or rejected before business records are created or updated.",
            ],
            [
                "One pipeline, multiple sources",
                "Text and voice should not create separate agent architectures.",
                "VOICE feeds transcript into the same AgentInstruction pipeline used by TEXT.",
            ],
            [
                "Reuse status standards",
                "Agent records use the same StatusScope and StatusUsage architecture as projects, risks, decisions, and risk actions.",
                "Agent-specific outcomes are modeled as fields and logs, not as narrow global statuses.",
            ],
            [
                "Clean business records",
                "Temporary operational behavior should not pollute final project records.",
                "WorkSession is separate from TimeEntry; only approved conversion creates the official TimeEntry.",
            ],
        ],
        [3.3, 6.4, 6.8],
    )

    add_heading(document, "3. Wave 4 Workstreams")
    add_table(
        document,
        ["Workstream", "Scope", "Outcome"],
        [
            [
                "Agent Core",
                "Instruction, suggestion, approval, action logging, source types, and status scopes.",
                "Reusable foundation for this and future agents.",
            ],
            [
                "Time Tracking Text Agent",
                "Start, pause, resume, finish, round duration, suggest time entry, approve into official TimeEntry.",
                "Fast operational time tracking without after-the-fact manual reconstruction.",
            ],
            [
                "Time Tracking Voice Pilot",
                "Capture voice, review/edit transcript, create AgentInstruction with sourceType VOICE.",
                "Early learning about voice quality, corrections, bilingual names, and command reliability.",
            ],
            [
                "Approval and Audit Trail",
                "Approval states, applied action logs, and traceability for all agent-created changes.",
                "A safe audit layer before multi-user and broader agents are introduced.",
            ],
            [
                "Configuration and Guardrails",
                "Visible agent behavior controls in the top-level Configuration area.",
                "No hidden seeded behavior; agent rules are understandable and adjustable before the assistant is used.",
            ],
        ],
        [4.2, 6.3, 6.0],
    )

    add_heading(document, "4. Data Model Contract")
    add_table(
        document,
        ["Model", "Role", "Key fields / relationships"],
        [
            [
                "AgentInstruction",
                "Stores the original user instruction or voice transcript.",
                "agentKey, sourceType, user, status, rawInstruction, transcript, normalizedInstruction, parsedIntentJson, optional project/workstream/task links.",
            ],
            [
                "AgentSuggestion",
                "Stores the structured proposal generated from an instruction.",
                "instruction, agentKey, suggestionType, targetEntity, targetRecordId, status, title, summary, payloadJson, appliedAt.",
            ],
            [
                "AgentApproval",
                "Stores the approval/rejection record for a suggestion.",
                "suggestion, approverUser, status, decisionNotes, decidedAt.",
            ],
            [
                "AgentActionLog",
                "Immutable event trail for agent activity.",
                "instruction, suggestion, approval, workSession, actorUser, agentKey, actionType, beforeJson, afterJson, metadataJson.",
            ],
            [
                "WorkSession",
                "Operational timer state before a TimeEntry exists.",
                "user, project, projectWorkstream, taskFamily, projectTask, status, sourceInstruction, convertedSuggestion, convertedTimeEntry, startedAt, endedAt, activeSeconds, roundedMinutes.",
            ],
            [
                "WorkSessionPause",
                "Pause/resume intervals for a WorkSession.",
                "workSession, pausedAt, resumedAt, notes.",
            ],
        ],
        [3.6, 4.8, 8.1],
    )

    add_heading(document, "5. Status Model")
    document.add_paragraph(
        "The agent foundation uses the existing Status, StatusScope, and StatusUsage architecture. This keeps lifecycle behavior consistent with the rest of the application and prepares the future security and translation layers."
    )
    add_table(
        document,
        ["Scope", "Default / lifecycle statuses", "Notes"],
        [
            [
                "AGENT_INSTRUCTION",
                "Open, In Progress, Closed, Cancelled",
                "Open is default. Represents whether an instruction is new, being processed, completed, or discarded.",
            ],
            [
                "AGENT_SUGGESTION",
                "Open, Approved, Rejected, Closed, Cancelled",
                "Open is default. Approved means the proposal passed human review; Closed can represent applied/completed state.",
            ],
            [
                "AGENT_APPROVAL",
                "Open, Approved, Rejected, Closed, Cancelled",
                "Open is default. The approval record is the human control point.",
            ],
            [
                "WORK_SESSION",
                "In Progress, On Hold, Closed, Cancelled",
                "In Progress is default. On Hold represents pause. Closed plus convertedTimeEntryId represents conversion without a special global status.",
            ],
        ],
        [4.2, 6.0, 6.3],
    )

    add_heading(document, "6. Time Tracking Assistant Behavior")
    add_numbered(
        document,
        [
            "User starts a work session by identifying project and workstream, optionally task family and task.",
            "The system creates or links an AgentInstruction and starts a WorkSession in the In Progress status.",
            "User can pause the session; the app records a WorkSessionPause and moves the session to On Hold.",
            "User can resume the session; the pause receives a resumedAt timestamp and the session returns to In Progress.",
            "User finishes the session; the app calculates active time excluding pauses and rounds to the configured increment.",
            "The app creates an AgentSuggestion to create a TimeEntry, not the TimeEntry itself.",
            "User approves, edits, or rejects the suggestion.",
            "Only after approval does the system create the official TimeEntry and link it back to the session.",
        ],
    )

    add_callout(
        document,
        "Rounding rule",
        "Initial rule: nearest 15-minute increment. Example: 1 hour 5 minutes rounds to 1.00 hour; 1 hour 8 minutes rounds to 1.25 hours. The rule is implemented as a reusable calculation and can later be made configurable.",
        fill=AMBER,
    )

    add_heading(document, "7. Voice Pilot Position")
    document.add_paragraph(
        "Voice should be tested earlier than originally planned, but it should not create a separate architecture. The voice pilot starts after the text instruction pipeline exists."
    )
    add_bullets(
        document,
        [
            "Voice is captured and converted to transcript.",
            "Transcript is shown to the user for review and correction.",
            "Confirmed transcript creates AgentInstruction with sourceType VOICE.",
            "The same suggestion, approval, action-log, and apply pipeline is reused.",
            "The pilot should explicitly test bilingual project/workstream names, noise, corrections, pause/resume commands, and ambiguity handling.",
        ],
    )

    add_heading(document, "8. Configuration and Guardrails")
    document.add_paragraph(
        "The agent foundation must include a visible configuration layer. Agent behavior should not depend on hidden seeded values or developer-only constants. The correct location is the top-level Configuration area, not Admin, because these settings determine how the application behaves. Admin remains focused on business reference data."
    )
    add_callout(
        document,
        "Build first",
        "Before building the Time Tracking Text Agent UI, create Configuration > Agents. This page should expose the agent's enabled state, source channels, approval rules, allowed capabilities, and time tracking rules. The assistant should read from this configuration instead of from hidden setup data.",
        fill=AMBER,
    )
    add_table(
        document,
        ["Configuration element", "Purpose", "Initial Wave 4 value"],
        [
            [
                "AgentDefinition",
                "Defines each available agent and whether it is enabled.",
                "Time Tracking Assistant enabled for one-user mode.",
            ],
            [
                "AgentCapability",
                "Defines what the agent is allowed to suggest or execute through approval.",
                "Start session, pause session, resume session, finish session, create time-entry suggestion.",
            ],
            [
                "AgentSourceConfig",
                "Controls input channels without changing the agent pipeline.",
                "TEXT enabled; VOICE defined but disabled until the pilot.",
            ],
            [
                "AgentRule",
                "Stores behavior rules that should be visible and controlled.",
                "Approval required, no auto-apply, one active session per user, 15-minute rounding.",
            ],
            [
                "Protected rules",
                "Rules that are safety constraints and should not be casually editable.",
                "No delete actions, no direct writes without approval, no changes to approved reports.",
            ],
        ],
        [4.0, 6.2, 6.3],
    )
    add_table(
        document,
        ["Area", "Configurable", "Protected in code"],
        [
            [
                "Input sources",
                "TEXT enabled first; VOICE enabled for pilot.",
                "Voice must still become transcript and instruction before any suggestion.",
            ],
            [
                "Rounding",
                "Increment and rounding mode can become configuration.",
                "Official TimeEntry is created only after approval.",
            ],
            [
                "Allowed targets",
                "Agent can be enabled for selected target entities.",
                "Agent cannot delete records or modify approved executive reports.",
            ],
            [
                "Approval",
                "Approval requirement can later vary by agent/action risk.",
                "Initial one-user assistant always requires approval before apply.",
            ],
            [
                "Session rules",
                "Cross-day handling, minimum duration, and one-active-session policy can be configured.",
                "Session must link to valid active project/workstream before conversion.",
            ],
        ],
        [3.4, 6.0, 7.1],
    )

    add_heading(document, "9. Implementation Waves")
    add_table(
        document,
        ["Wave", "Build", "Stop / review checkpoint"],
        [
            [
                "4.0 Foundation",
                "Data model, status scopes, typed helpers, action log service, work-session duration utilities, and visible Configuration > Agents behavior controls.",
                "Schema validates, migration applies, build/lint pass, and agent behavior is visible before operational UI is added.",
            ],
            [
                "4.0A Agent Configuration",
                "Create top-level Configuration > Agents with enabled agents, source channels, capabilities, approval rules, rounding, and session constraints.",
                "Manual test confirms behavior settings can be reviewed and changed without using seed scripts.",
            ],
            [
                "4.1 Text Time Tracking Assistant",
                "Marisa-only session UI: start, pause, resume, finish, create suggestion, approval into TimeEntry.",
                "Manual test confirms the assistant reads configuration and follows session lifecycle, rounding, approval, and official TimeEntry creation rules.",
            ],
            [
                "4.2 Voice Pilot",
                "Voice capture/transcript review creates AgentInstruction with sourceType VOICE.",
                "Manual test confirms transcript correction and reuse of the same approval pipeline.",
            ],
            [
                "Later expansion",
                "Progress reporting agent, broader security model, multi-user agents, and wider audit logging.",
                "Do only after one-user time tracking assistant proves useful and understandable.",
            ],
        ],
        [3.1, 7.0, 6.4],
    )

    add_heading(document, "10. Test Criteria")
    add_bullets(
        document,
        [
            "A text instruction can be stored as AgentInstruction without creating business data.",
            "Agent configuration can be reviewed in Configuration > Agents before any assistant behavior is used.",
            "The Time Tracking Assistant reads enabled sources, approval requirement, rounding increment, and session constraints from configuration.",
            "A WorkSession can start, pause, resume, finish, and calculate active duration excluding pauses.",
            "Rounding follows the agreed 15-minute nearest-increment rule.",
            "A finished session creates an AgentSuggestion, not a TimeEntry.",
            "A TimeEntry is created only after explicit approval.",
            "All lifecycle events create AgentActionLog records.",
            "Rejected or cancelled suggestions do not alter business data.",
            "Voice transcript review can later feed the same instruction pipeline without new agent tables.",
        ],
    )

    add_heading(document, "11. Current Implementation Note")
    add_callout(
        document,
        "Foundation status",
        "At the time of this document, the structural Wave 4.0 foundation has already been created in the codebase: Prisma models, migration, status scopes/usages, and typed domain helpers. The next foundation step is to add visible Agent Configuration so behavior is not controlled by hidden seed/setup data.",
        fill=GREEN,
    )

    add_heading(document, "12. Recommended Next Step")
    document.add_paragraph(
        "Review and confirm this blueprint. Once confirmed, continue with Wave 4.0A by building the top-level Configuration > Agents page and backing configuration model. Only after that should Wave 4.1 build the Marisa-only Time Tracking Text Agent."
    )

    add_footer(document)
    document.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
