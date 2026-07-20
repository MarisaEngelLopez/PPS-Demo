from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "PPS_Production_Recovery_Runbook_2026-07-18.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Code Block"
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)


def add_kv_table(doc, rows, widths=(2700, 6660)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, list(widths))
    first = table.rows[0]
    first.cells[0].text = rows[0][0]
    first.cells[1].text = rows[0][1]
    for cell in first.cells:
        set_cell_shading(cell, "F2F4F7")
    for label, value in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
            if idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 58, 95)
    title.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after in [
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181) if style_name != "Heading 3" else RGBColor(31, 77, 120)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="DADCE0")
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(f" {body}")
    return table


def build():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("PPS Production Recovery Runbook")

    subtitle = doc.add_paragraph()
    subtitle.add_run("Frozen code recovery with current live database preserved").bold = True
    subtitle.add_run(" | Validated 2026-07-18")

    add_note(
        doc,
        "Outcome:",
        "Production was recovered from the frozen v1.1.0 code backup while preserving the current live dev.db. Records created after the backup remained visible after recovery.",
    )

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This runbook documents the validated recovery path for the Project Operations System production environment after development code was accidentally run against the live database."
    )
    add_bullets(
        doc,
        [
            "Recover production application code from a frozen release backup.",
            "Preserve the current live database, including records created after the code backup.",
            "Restore runtime configuration that is intentionally excluded from source backups.",
            "Generate Prisma Client inside the recovered code folder before validating menu/data pages.",
        ],
    )

    doc.add_heading("2. Recovery Architecture", level=1)
    add_kv_table(
        doc,
        [
            ("Item", "Value"),
            ("Production code folder", r"C:\Users\maris\Documents\Project Management system 2\project-ops-system-production"),
            ("Development code folder", r"C:\Users\maris\Documents\Project Management system 2\project-ops-system"),
            ("Live database", r"C:\Users\maris\Documents\Project Management system 2\dev.db"),
            ("Production port", "http://localhost:3000"),
            ("Development port", "http://localhost:3001"),
            ("Recovery code source", r"backups\release_v1.1.0_20260715_180907\project-ops-system-v1.1.0-code.zip"),
            ("Runtime config source", r"project-ops-system\.env.local"),
        ],
    )

    doc.add_heading("3. Incident Summary", level=1)
    doc.add_paragraph(
        "During V3.3 development, the active codebase gained new Prisma schema expectations such as workspace scope columns and auth tables. Production was still using the current live database, which intentionally had not been migrated. Running the V3.3 code on production caused server errors such as:"
    )
    add_code_block(doc, "The column main.Organization.workspaceId does not exist in the current database.")
    doc.add_paragraph(
        "The correct recovery was not to restore the older database backup, because real production work had continued after the backup. The correct recovery was to restore frozen production code and keep the current live database."
    )

    doc.add_heading("4. Recovery Procedure", level=1)
    doc.add_paragraph("Run these commands from PowerShell unless otherwise stated.")

    doc.add_heading("4.1 Stop Production", level=2)
    doc.add_paragraph("In the production PowerShell window, stop the running app:")
    add_code_block(doc, "Ctrl+C")

    doc.add_heading("4.2 Create or Refresh the Frozen Production Folder", level=2)
    doc.add_paragraph("If the production folder does not exist, expand the frozen code backup:")
    add_code_block(
        doc,
        'cd "C:\\Users\\maris\\Documents\\Project Management system 2"\n'
        'Expand-Archive -Path "backups\\release_v1.1.0_20260715_180907\\project-ops-system-v1.1.0-code.zip" -DestinationPath "project-ops-system-production"',
    )
    doc.add_paragraph(
        "If the production folder already exists but a file was accidentally changed, restore that file from the same backup zip instead of editing frozen code manually."
    )

    doc.add_heading("4.3 Restore Runtime Configuration", level=2)
    doc.add_paragraph(
        ".env.local is intentionally excluded from the source backup because it can contain private runtime values. Copy it from the current development folder into the frozen production folder:"
    )
    add_code_block(
        doc,
        'Copy-Item "project-ops-system\\.env.local" "project-ops-system-production\\.env.local" -Force',
    )

    doc.add_heading("4.4 Install Dependencies", level=2)
    doc.add_paragraph(
        "If node_modules is missing in the recovered production folder, install dependencies from the lockfile. Do not run npm audit fix during recovery because it changes the frozen dependency set."
    )
    add_code_block(
        doc,
        'cd "C:\\Users\\maris\\Documents\\Project Management system 2\\project-ops-system-production"\n'
        "$env:NODE_OPTIONS='--use-system-ca'\n"
        "npm ci",
    )

    doc.add_heading("4.5 Generate Prisma Client", level=2)
    doc.add_paragraph(
        "This is required after restoring code and dependencies. It writes generated client files under node_modules only. It does not migrate or write to the live database."
    )
    add_code_block(
        doc,
        "$env:DATABASE_URL='file:../dev.db'\n"
        "npx prisma generate",
    )

    doc.add_heading("4.6 Validate Prisma Schema Without Writing to DB", level=2)
    add_code_block(
        doc,
        "$env:DATABASE_URL='file:../dev.db'\n"
        "npx prisma validate",
    )

    doc.add_heading("4.7 Start Production", level=2)
    doc.add_paragraph("Start production from the parent folder using the launcher:")
    add_code_block(
        doc,
        'cd "C:\\Users\\maris\\Documents\\Project Management system 2"\n'
        ".\\Start-PPS-Production.ps1",
    )

    doc.add_heading("5. Validation Checklist", level=1)
    add_bullets(
        doc,
        [
            "Open http://localhost:3000.",
            "Navigate through menu items that load server data.",
            "Confirm records created after the backup are visible.",
            "Confirm the missing workspaceId error no longer appears.",
            "Confirm the app is using the current live dev.db, not the backup database snapshot.",
            "Keep V3.3 development testing on http://localhost:3001 only.",
        ],
    )

    doc.add_heading("6. Important Rules", level=1)
    add_bullets(
        doc,
        [
            "Do not restore dev-v1.1.0.db over the current live dev.db if production work has continued.",
            "Do not migrate the live database as part of this recovery unless a separate approved migration plan exists.",
            "Do not edit frozen production code during recovery; restore from the backup source instead.",
            "Do not run npm audit fix during recovery; it can change the frozen dependency baseline.",
            "Always generate Prisma Client after dependency installation in a recovered folder.",
        ],
    )

    doc.add_heading("7. Launch Commands", level=1)
    add_kv_table(
        doc,
        [
            ("Environment", "Command"),
            ("Production", r'cd "C:\Users\maris\Documents\Project Management system 2"; .\Start-PPS-Production.ps1'),
            ("Development", r'cd "C:\Users\maris\Documents\Project Management system 2"; .\Start-PPS-Development.ps1'),
        ],
        widths=(2000, 7360),
    )

    doc.add_heading("8. Recovery Evidence Captured", level=1)
    add_bullets(
        doc,
        [
            "Production opened successfully after recovery.",
            "Menu navigation worked after Prisma Client generation.",
            "Records created after the backup were visible, proving the current live database was preserved.",
            "The production code path is now separated from the V3.3 development code path.",
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Document owner: Project Operations System recovery runbook").italic = True

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
